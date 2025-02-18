import streamlit as st
from crewai import Agent, Task, Crew
import os
from crewai_tools import SerperDevTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64

load_dotenv()

st.set_page_config(layout="wide")
st.markdown("""
    <style>
        .title { font-weight: bold; color: #191970; font-size: 24px; }
        .box {
            background-color: #104e8b;
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 15px;
        }
        .button {
            background-color: #191970;
            color: white;
            font-weight: bold;
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            cursor: pointer;
            width: 100%;
        }
    </style>
""", unsafe_allow_html=True)

st.title("AI Medical Assistant")

col1, col2 = st.columns([1, 2])

# **Sol tarafta API Seçimi**
with col1:
    st.markdown("<div class='box'>", unsafe_allow_html=True)
    st.markdown("<p class='title'>Choose API Type</p>", unsafe_allow_html=True)
    api_type = st.selectbox("", ("OpenAI API", "Google API"))
    st.markdown(f"<p class='title'>Enter your {api_type} Key</p>", unsafe_allow_html=True)
    api_key = st.text_input("", type="password")
    st.markdown("</div>", unsafe_allow_html=True)

# **Sağ tarafta Kullanıcıdan Medikal Bilgi Alımı**
with col2:
    st.markdown("<div class='box'>", unsafe_allow_html=True)
    st.markdown("<p class='title'>Medical Information</p>", unsafe_allow_html=True)
    gender = st.selectbox('Gender', ('Male', 'Female', 'Other'))
    age = st.number_input('Age', min_value=0, max_value=120, value=25)
    symptoms = st.text_area('Symptoms', 'e.g., fever, cough')
    medical_history = st.text_area('Medical History', 'e.g., diabetes')
    st.markdown("</div>", unsafe_allow_html=True)

# **Dosya Oluşturma Fonksiyonları**
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment', 0)
    doc.add_paragraph(str(result))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Report</a>'

# **API Anahtarını Ayarlama**
if api_key:
    if api_type == "OpenAI API":
        os.environ["OPENAI_API_KEY"] = api_key
    elif api_type == "Google API":
        os.environ["GOOGLE_API_KEY"] = api_key

# **CrewAI ve LLM Modeli Tanımlama**
search_tool = SerperDevTool()
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0,
    max_tokens=3000
)

diagnostician = Agent(
    role="Diagnostician",
    goal="Identify possible conditions from symptoms.",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history.",
    verbose=False,
    allow_delegation=False,
    tools=[search_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Suggest effective treatments.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs.",
    verbose=False,
    allow_delegation=False,
    tools=[search_tool],
    llm=llm
)

# **Crew ve Task Tanımları**
diagnose_task = Task(
    description=f"Analyze symptoms: {symptoms} and history: {medical_history}. Provide possible conditions.",
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician
)

treatment_task = Task(
    description=f"Based on diagnosis, recommend treatment considering history: {medical_history}.",
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor
)

crew = Crew(agents=[diagnostician, treatment_advisor], tasks=[diagnose_task, treatment_task], verbose=False)

# **Butonun Doğru Çalışması İçin Güncellendi**
with col1:
    if st.button("Get Diagnosis & Treatment", key="start", help="Click to start diagnosis and treatment", use_container_width=True):
        if not api_key:
            st.error("Please enter an API key before proceeding.")
        else:
            with st.spinner('Processing...'):
                result = crew.kickoff(inputs={"symptoms": str(symptoms), "medical_history": str(medical_history)})
                st.write(result)
                docx_file = generate_docx(result)
                st.markdown(get_download_link(docx_file, "diagnosis_report.docx"), unsafe_allow_html=True)
